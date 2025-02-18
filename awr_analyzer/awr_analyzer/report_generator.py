import pandas as pd
from docx import Document

def generate_awr_report(parsed_awr, output_path):
    """Gera um relatório DOCX detalhado do AWR analisado, incluindo explicações e recomendações baseadas nos dados reais extraídos."""
    print("✅ Gerando o relatório...")
    document = Document()
    document.add_heading("AWR Analysis Report", 0)
    
    # 🔍 CPU Usage
    document.add_heading("CPU Usage", level=1)
    cpu_data = parsed_awr.get("CPU Usage", [])
    if cpu_data and cpu_data[0] != "Seção não encontrada":
        for item in cpu_data:
            document.add_paragraph(f"🔹 Uso de CPU: {item}", style="List Bullet")
        document.add_paragraph("⚠️ Se o uso de CPU estiver alto, investigue SQLs pesados e revise estatísticas de índices.")
    else:
        document.add_paragraph("⚠️ Dados de CPU não encontrados no relatório AWR.")
    
    # 🔍 PGA Usage
    document.add_heading("PGA Usage", level=1)
    pga_data = parsed_awr.get("PGA Usage", [])
    if pga_data and pga_data[0] != "Seção não encontrada":
        for item in pga_data:
            document.add_paragraph(f"🔹 {item[0]}: {item[1]}", style="List Bullet")
        document.add_paragraph("⚠️ Se o uso de PGA estiver alto, ajuste `PGA_AGGREGATE_TARGET` e revise operações que utilizam alta memória.")
    else:
        document.add_paragraph("⚠️ Dados de PGA não encontrados no relatório AWR.")
    
    # 🔍 Wait Events
    document.add_heading("Wait Events", level=1)
    wait_events = parsed_awr.get("Wait Events", [])
    if wait_events and wait_events[0][0] != "Seção não encontrada":
        for event in wait_events:
            document.add_paragraph(f"🔹 Evento: {event[0]}, Tempo total: {event[1]}, Tipo: {event[2]}", style="List Bullet")
        document.add_paragraph("⚠️ Investigue se os eventos de espera indicam contenção de I/O, locks ou problemas de CPU.")
    else:
        document.add_paragraph("⚠️ Nenhum evento de espera encontrado.")
    
    # 🔍 SQL Offenders
    document.add_heading("SQL Offenders", level=1)
    sql_offenders = parsed_awr.get("SQL Offenders", [])
    if sql_offenders and sql_offenders[0][0] != "Seção não encontrada":
        for sql in sql_offenders:
            document.add_paragraph(f"🔍 SQL ID: {sql[6]}", style="List Bullet")
            document.add_paragraph(f"👉 Query: {sql[8]}")
            document.add_paragraph(f"⚠️ Tempo de execução: {sql[0]}s, Execuções: {sql[1]}")
            document.add_paragraph("✅ Recomendações:")
            document.add_paragraph("- Utilize `EXPLAIN PLAN` e `DBMS_XPLAN.DISPLAY_CURSOR` para analisar gargalos.", style="List Bullet")
            document.add_paragraph("- Verifique índices e estatísticas de tabelas (`DBMS_STATS.GATHER_TABLE_STATS`).", style="List Bullet")
            document.add_paragraph("- Avalie possíveis reescritas de queries para otimizar os planos de execução.", style="List Bullet")
    else:
        document.add_paragraph("⚠️ Nenhum SQL crítico identificado.")
    
    # 🔍 ADDM Findings
    document.add_heading("ADDM Findings", level=1)
    addm_findings = parsed_awr.get("ADDM Findings", [])
    if addm_findings and addm_findings[0][0] != "Seção não encontrada":
        for finding in addm_findings:
            document.add_paragraph(f"🔹 Relatório: {finding[0]}, Impacto: {finding[1]}", style="List Bullet")
        document.add_paragraph("⚠️ Utilize `DBMS_ADVISOR.TUNE_SQL` para recomendações detalhadas.")
    else:
        document.add_paragraph("⚠️ Nenhuma recomendação ADDM encontrada.")
    
    # 📌 Plano de Ação
    document.add_heading("Plano de Ação", level=1)
    document.add_paragraph("Baseado na análise, siga estas recomendações:")
    document.add_paragraph("1️⃣ Revisar e otimizar queries ofensivas utilizando índices e estatísticas atualizadas.", style="List Number")
    document.add_paragraph("2️⃣ Ajustar `SGA` e `PGA` conforme necessidade para evitar contenção de memória.", style="List Number")
    document.add_paragraph("3️⃣ Identificar eventos de espera e aplicar ajustes para reduzir bloqueios e contenção de I/O.", style="List Number")
    document.add_paragraph("4️⃣ Revisar parâmetros de paralelismo (`PARALLEL_EXECUTION_MESSAGE_SIZE`, `CPU_COUNT`) para otimizar workload.", style="List Number")
    
    # Salvando o relatório
    document.save(output_path)
    print(f"✅ Relatório gerado com sucesso em: {output_path}")

    # 📚 Links da Documentação
    document.add_heading("Oracle Documentation Links", level=1)
    oracle_docs = {
        "SQL Tuning": "https://docs.oracle.com/en/database/oracle/oracle-database/19/tgsql/",
        "Index Optimization": "https://docs.oracle.com/en/database/oracle/oracle-database/19/admin/managing-indexes.html",
        "Parallel Execution": "https://docs.oracle.com/en/database/oracle/oracle-database/19/dwhsg/parallel-execution.html"
    }
    for topic, link in oracle_docs.items():
        document.add_paragraph(f"{topic}: {link}", style="List Bullet")

    document.save(output_path)
    print(f"✅ Relatório gerado com sucesso em: {output_path}")
