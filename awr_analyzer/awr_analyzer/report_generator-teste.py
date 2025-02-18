import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
import time

from awr_analyzer.sql_analyzer import analyze_sql_ordered_by_elapsed_time, generate_sql_recommendations
from awr_analyzer.awr_offenders import extract_offenders
from awr_analyzer.oracle_docs import get_oracle_doc


def generate_awr_report(awr_file_path, output_path):
    """
    Gera um relatório DOCX detalhado do AWR analisado.
    Inclui análise de SQLs lentos, CPU, I/O, Locks, Memória e Paralelismo.
    """
    print("✅ Gerando o relatório...")

    # Parse AWR file
    try:
        with open(awr_file_path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file, 'html.parser')
    except UnicodeDecodeError:
        with open(awr_file_path, 'r', encoding='latin-1') as file:
            soup = BeautifulSoup(file, 'html.parser')

    document = Document()
    document.add_heading("AWR Analysis Report", 0)

    # 🔍 Analisando SQLs problemáticos
    document.add_heading("SQL Performance Analysis", level=1)
    sql_df = analyze_sql_ordered_by_elapsed_time(soup)
    if sql_df is not None:
        sql_recommendations = generate_sql_recommendations(sql_df)
        for rec in sql_recommendations:
            document.add_paragraph(rec, style="List Bullet")

    # 🔎 Analisando ofensores do AWR
    offenders = extract_offenders(soup)
    for section, data in offenders.items():
        document.add_heading(section, level=1)
        for item in data:
            document.add_paragraph(item, style="List Bullet")

    # 📌 Plano de Ação Detalhado
    document.add_heading("Plano de Ação", level=1)
    
    actions = {
        "🔥 Uso de CPU": offenders.get("CPU Usage", []),
        "💾 I/O Excessivo": offenders.get("IO Profile", []),
        "🔒 Locks e Contenção": offenders.get("Lock Contention", []),
        "⚡ Parallelismo": offenders.get("Parallel Execution", []),
        "🧠 Memória (SGA/PGA)": offenders.get("Memory Usage", [])
    }

    for section, details in actions.items():
        if details:
            document.add_heading(section, level=2)
            for issue in details:
                document.add_paragraph(issue, style="List Bullet")

    # 📚 Links para Documentação Oficial da Oracle
    document.add_heading("Oracle Documentation Links", level=1)
    oracle_docs = {
        "SQL Tuning": get_oracle_doc("sql_tuning"),
        "Index Optimization": get_oracle_doc("index_tuning"),
        "Parallel Execution": get_oracle_doc("parallelism"),
        "Memory Tuning": get_oracle_doc("memory_tuning"),
        "Locks & Contention": get_oracle_doc("locks_contention"),
        "I/O Performance": get_oracle_doc("io_performance")
    }

    for topic, link in oracle_docs.items():
        document.add_paragraph(f"{topic}: {link}", style="List Bullet")

    # 💾 Tratamento de erro ao salvar o arquivo
    save_document(document, output_path)


def save_document(document, output_path):
    """
    Tenta salvar o documento, detectando se o arquivo está aberto e pedindo ao usuário para fechá-lo antes de continuar.
    """
    max_attempts = 3  # Número de tentativas antes de falhar
    attempt = 0

    while attempt < max_attempts:
        try:
            document.save(output_path)
            print(f"✅ Relatório gerado com sucesso em: {output_path}")
            return
        except PermissionError:
            print(f"⚠️ O arquivo '{output_path}' está aberto. Por favor, feche-o e pressione Enter para tentar novamente...")
            input("Pressione Enter após fechar o arquivo...")
            attempt += 1

    print("❌ Não foi possível salvar o arquivo. Verifique se ele não está sendo usado e tente novamente.")


if __name__ == "__main__":
    awr_file = "C:/temp/awr_analyzer/input/awr_report.html"
    output_file = "C:/temp/awr_analyzer/output/awr_analysis_report.docx"
    generate_awr_report(awr_file, output_file)
