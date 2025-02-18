import pandas as pd
from docx import Document

from bs4 import BeautifulSoup

def parse_awr_report(awr_html):
    """
    Faz o parsing do relatório AWR em HTML e extrai os principais dados.
    Retorna um dicionário com as informações extraídas.
    """
    soup = BeautifulSoup(awr_html, 'html.parser')
    data = {}

    def extract_table_data(title_variations):
        """
        Procura uma seção pelo título e retorna os dados da tabela seguinte.
        Busca nos elementos <h1>, <h2>, <h3>, <th>, <b>, <font>.
        """
        section = None
        for title in title_variations:
            section = soup.find(["h1", "h2", "h3", "th", "b", "font"], string=lambda text: text and title.lower() in text.lower())
            if section:
                break  # Para ao encontrar o título correto
        
        if section:
            table = section.find_next("table")
            if table:
                rows = table.find_all("tr")[1:]  # Ignorar cabeçalhos
                return [[col.text.strip() for col in row.find_all("td")] for row in rows]
        return ["Seção não encontrada"]

    # 🔍 Substitua os títulos abaixo pelos que apareceram no **Passo 1**
    data["CPU Usage"] = extract_table_data(["Host CPU Usage", "Host CPU Statistics", "CPU Utilization"])
    data["PGA Usage"] = extract_table_data(["PGA Usage Summary", "PGA Memory Usage", "PGA Statistics"])
    data["Wait Events"] = extract_table_data(["Top Timed Events", "Wait Events", "Performance Events"])
    data["SQL Offenders"] = extract_table_data(["SQL Ordered by Elapsed Time", "SQL Top Queries", "SQL Execution Time"])
    data["ADDM Findings"] = extract_table_data(["ADDM Findings", "Performance Recommendations", "Automatic Database Diagnostic Monitor Findings"])

    return data

    
    # 📌 Plano de Ação
    document.add_heading("Plano de Ação", level=1)
    document.add_paragraph("Baseado na análise, siga estas recomendações:")
    document.add_paragraph("1️⃣ Revisar e otimizar queries ofensivas utilizando índices e estatísticas atualizadas.", style="List Number")
    document.add_paragraph("2️⃣ Ajustar `SGA` e `PGA` conforme necessidade para evitar contenção de memória.", style="List Number")
    document.add_paragraph("3️⃣ Identificar eventos de espera e aplicar ajustes para reduzir bloqueios e contenção de I/O.", style="List Number")
    document.add_paragraph("4️⃣ Revisar parâmetros de paralelismo (`PARALLEL_EXECUTION_MESSAGE_SIZE`, `CPU_COUNT`) para otimizar workload.", style="List Number")
    
    # 📚 Documentação Oficial da Oracle
    document.add_heading("Oracle Documentation Links", level=1)
    oracle_docs = {
        "SQL Tuning": "https://docs.oracle.com/en/database/oracle/oracle-database/19/tgsql/",
        "Index Optimization": "https://docs.oracle.com/en/database/oracle/oracle-database/19/admin/managing-indexes.html",
        "Parallel Execution": "https://docs.oracle.com/en/database/oracle/oracle-database/19/dwhsg/parallel-execution.html",
        "Memory Tuning": "https://docs.oracle.com/en/database/oracle/oracle-database/19/tgsql/tuning-memory.html",
        "Locks & Contention": "https://docs.oracle.com/en/database/oracle/",
        "I/O Performance": "https://docs.oracle.com/en/database/oracle/"
    }
    for topic, link in oracle_docs.items():
        document.add_paragraph(f"{topic}: {link}", style="List Bullet")
    
    # Salvando o relatório
    document.save(output_path)
    print(f"✅ Relatório gerado com sucesso em: {output_path}")
