import os
import yt_dlp
import whisper
from docx import Document
import warnings
warnings.simplefilter("ignore")

def download_audio(youtube_url, output_folder="downloads"):
    """Baixa o áudio do vídeo do YouTube."""
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    ydl_opts = {
        'format': 'bestaudio/best',
        'outtmpl': f'{output_folder}/audio.%(ext)s',
        'postprocessors': [{
            'key': 'FFmpegExtractAudio',
            'preferredcodec': 'mp3',
            'preferredquality': '192',
        }],
    }
    
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        ydl.download([youtube_url])
    
    return f"{output_folder}/audio.mp3"

def transcribe_audio(audio_file):
    """Transcreve o áudio usando Whisper."""
    model = whisper.load_model("base")
    result = model.transcribe(audio_file)
    return result["text"]

def save_transcription(transcription, output_file="transcription.docx"):
    """Salva a transcrição em um arquivo Word (.docx)."""
    doc = Document()
    doc.add_heading("Transcrição do Vídeo", level=1)
    doc.add_paragraph(transcription)
    doc.save(output_file)
    return output_file

def main():
    youtube_url = input("Digite o link do vídeo do YouTube: ")
    print("Baixando o áudio...")
    audio_file = download_audio(youtube_url)
    
    print("Transcrevendo o áudio...")
    transcription = transcribe_audio(audio_file)
    
    print("Salvando transcrição...")
    output_file = save_transcription(transcription)
    
    print(f"Transcrição salva em: {output_file}")
    
if __name__ == "__main__":
    main()